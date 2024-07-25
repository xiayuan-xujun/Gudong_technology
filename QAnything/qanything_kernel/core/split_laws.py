from langchain.docstore.document import Document
from docx import Document as DocumentDocx
import re

BULLET_PATTERN = [[
    r"第[零一二三四五六七八九十百0-9]+(分?编|部分)",
    r"第[零一二三四五六七八九十百0-9]+章",
    r"第[零一二三四五六七八九十百0-9]+节",
    r"第[零一二三四五六七八九十百0-9]+条",
    r"[\(（][零一二三四五六七八九十百]+[\)）]",
], [
    r"第[0-9]+章",
    r"第[0-9]+节",
    r"[0-9]{,2}[\. 、]",
    r"[0-9]{,2}\.[0-9]{,2}[^a-zA-Z/%~-]",
    r"[0-9]{,2}\.[0-9]{,2}\.[0-9]{,2}",
    r"[0-9]{,2}\.[0-9]{,2}\.[0-9]{,2}\.[0-9]{,2}",
], [
    r"第[零一二三四五六七八九十百0-9]+章",
    r"第[零一二三四五六七八九十百0-9]+节",
    r"[零一二三四五六七八九十百]+[ 、]",
    r"[\(（][零一二三四五六七八九十百]+[\)）]",
    r"[\(（][0-9]{,2}[\)）]",
], [
    r"PART (ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN)",
    r"Chapter (I+V?|VI*|XI|IX|X)",
    r"Section [0-9]+",
    r"Article [0-9]+"
]
]

def not_bullet(line):
    patt = [
        r"0", r"[0-9]+ +[0-9~个只-]", r"[0-9]+\.{2,}"
    ]
    return any([re.match(r, line) for r in patt])


def bullets_category(sections):
    global BULLET_PATTERN
    hits = [0] * len(BULLET_PATTERN)
    for i, pro in enumerate(BULLET_PATTERN):
        for sec in sections:
            for p in pro:
                if re.match(p, sec) and not not_bullet(sec):
                    hits[i] += 1
                    break
    maxium = 0
    res = -1
    for i, h in enumerate(hits):
        if h <= maxium:
            continue
        res = i
        maxium = h
    return res


def docx_question_level(p, bull = -1):
    txt = re.sub(r"\u3000", " ", p.text).strip()
    if p.style.name.startswith('Heading'):
        return int(p.style.name.split(' ')[-1]), txt
    else:
        if bull < 0:
            return 0, txt
        for j, title in enumerate(BULLET_PATTERN[bull]):
            if re.match(title, txt):
                return j+1, txt
    return len(BULLET_PATTERN[bull]), txt


def split_doc(file_path):
    doc = DocumentDocx(file_path)
    pn = 0
    lines = []
    bull = bullets_category([p.text for p in doc.paragraphs])
    for p in doc.paragraphs:
        if pn > 100000:
            break
        question_level, p_text = docx_question_level(p, bull)
        if not p_text.strip("\n"):
            continue
        lines.append((question_level, p_text))

        for run in p.runs:
            if 'lastRenderedPageBreak' in run._element.xml:
                pn += 1
                continue
            if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                pn += 1

    visit = [False for _ in range(len(lines))]
    sections = []
    for s in range(len(lines)):
        e = s + 1
        while e < len(lines):
            if lines[e][0] <= lines[s][0]:
                break
            e += 1
        if e - s == 1 and visit[s]:
            continue
        sec = []
        next_level = lines[s][0] + 1
        while not sec and next_level < 22:
            for i in range(s+1, e):
                if lines[i][0] != next_level:
                    continue
                sec.append(lines[i][1])
                visit[i] = True
            next_level += 1
        sec.insert(0, lines[s][1])

        sections.append("\n".join(sec))
    return [l for l in sections if l]

if __name__ == '__main__':
    file_path = r'/home/deeplearn/workspace/wxj/QAnything_old/docs/docx/测试数据.docx'
    docs = split_doc(file_path)
    # 转成doc格式的列表数据
    new_docs = []
    for idx, doc in enumerate(docs):
        page_content = re.sub(r'[\n\t]+', '\n', doc).strip()
        new_doc = Document(page_content=page_content)
        new_doc.metadata["user_id"] = "self.user_id"
        new_doc.metadata["kb_id"] = "self.kb_id"
        new_doc.metadata["file_id"] = "self.file_id"
        new_doc.metadata["file_name"] = "self.url if self.url else self.file_name"
        new_doc.metadata["chunk_id"] = idx
        new_doc.metadata["file_path"] = "self.file_path"
        new_doc.metadata['faq_dict'] = {}
        new_docs.append(new_doc)