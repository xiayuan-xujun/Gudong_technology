#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

from docx import Document
import re
import pandas as pd
from collections import Counter
from rag.nlp import rag_tokenizer
from io import BytesIO

def is_chinese(s):
    if s >= u'\u4e00' and s <= u'\u9fa5':
        return True
    else:
        return False


def is_number(s):
    if s >= u'\u0030' and s <= u'\u0039':
        return True
    else:
        return False


def is_alphabet(s):
    if (s >= u'\u0041' and s <= u'\u005a') or (
            s >= u'\u0061' and s <= u'\u007a'):
        return True
    else:
        return False


def naiveQie(txt):
    tks = []
    for t in txt.split(" "):
        if tks and re.match(r".*[a-zA-Z]$", tks[-1]
                            ) and re.match(r".*[a-zA-Z]$", t):
            tks.append(" ")
        tks.append(t)
    return tks

class RagTokenizer:
    def key_(self, line):
        return str(line.lower().encode("utf-8"))[2:-1]

    def rkey_(self, line):
        return str(("DD" + (line[::-1].lower())).encode("utf-8"))[2:-1]

    def loadDict_(self, fnm):
        print("[HUQIE]:Build trie", fnm, file=sys.stderr)
        try:
            of = open(fnm, "r", encoding='utf-8')
            while True:
                line = of.readline()
                if not line:
                    break
                line = re.sub(r"[\r\n]+", "", line)
                line = re.split(r"[ \t]", line)
                k = self.key_(line[0])
                F = int(math.log(float(line[1]) / self.DENOMINATOR) + .5)
                if k not in self.trie_ or self.trie_[k][0] < F:
                    self.trie_[self.key_(line[0])] = (F, line[2])
                self.trie_[self.rkey_(line[0])] = 1
            self.trie_.save(fnm + ".trie")
            of.close()
        except Exception as e:
            print("[HUQIE]:Faild to build trie, ", fnm, e, file=sys.stderr)

    def __init__(self, debug=False):
        self.DEBUG = debug
        self.DENOMINATOR = 1000000
        self.trie_ = datrie.Trie(string.printable)
        self.DIR_ = os.path.join(get_project_base_directory(), "rag/res", "huqie")

        self.stemmer = PorterStemmer()
        self.lemmatizer = WordNetLemmatizer()

        self.SPLIT_CHAR = r"([ ,\.<>/?;'\[\]\\`!@#$%^&*\(\)\{\}\|_+=《》，。？、；‘’：“”【】~！￥%……（）——-]+|[a-z\.-]+|[0-9,\.-]+)"
        try:
            self.trie_ = datrie.Trie.load(self.DIR_ + ".txt.trie")
            return
        except Exception as e:
            print("[HUQIE]:Build default trie", file=sys.stderr)
            self.trie_ = datrie.Trie(string.printable)

        self.loadDict_(self.DIR_ + ".txt")

    def loadUserDict(self, fnm):
        try:
            self.trie_ = datrie.Trie.load(fnm + ".trie")
            return
        except Exception as e:
            self.trie_ = datrie.Trie(string.printable)
        self.loadDict_(fnm)

    def addUserDict(self, fnm):
        self.loadDict_(fnm)

    def _strQ2B(self, ustring):
        """把字符串全角转半角"""
        rstring = ""
        for uchar in ustring:
            inside_code = ord(uchar)
            if inside_code == 0x3000:
                inside_code = 0x0020
            else:
                inside_code -= 0xfee0
            if inside_code < 0x0020 or inside_code > 0x7e:  # 转完之后不是半角字符返回原来的字符
                rstring += uchar
            else:
                rstring += chr(inside_code)
        return rstring

    def _tradi2simp(self, line):
        return HanziConv.toSimplified(line)

    def dfs_(self, chars, s, preTks, tkslist):
        MAX_L = 10
        res = s
        # if s > MAX_L or s>= len(chars):
        if s >= len(chars):
            tkslist.append(preTks)
            return res

        # pruning
        S = s + 1
        if s + 2 <= len(chars):
            t1, t2 = "".join(chars[s:s + 1]), "".join(chars[s:s + 2])
            if self.trie_.has_keys_with_prefix(self.key_(t1)) and not self.trie_.has_keys_with_prefix(
                    self.key_(t2)):
                S = s + 2
        if len(preTks) > 2 and len(
                preTks[-1][0]) == 1 and len(preTks[-2][0]) == 1 and len(preTks[-3][0]) == 1:
            t1 = preTks[-1][0] + "".join(chars[s:s + 1])
            if self.trie_.has_keys_with_prefix(self.key_(t1)):
                S = s + 2

        ################
        for e in range(S, len(chars) + 1):
            t = "".join(chars[s:e])
            k = self.key_(t)

            if e > s + 1 and not self.trie_.has_keys_with_prefix(k):
                break

            if k in self.trie_:
                pretks = copy.deepcopy(preTks)
                if k in self.trie_:
                    pretks.append((t, self.trie_[k]))
                else:
                    pretks.append((t, (-12, '')))
                res = max(res, self.dfs_(chars, e, pretks, tkslist))

        if res > s:
            return res

        t = "".join(chars[s:s + 1])
        k = self.key_(t)
        if k in self.trie_:
            preTks.append((t, self.trie_[k]))
        else:
            preTks.append((t, (-12, '')))

        return self.dfs_(chars, s + 1, preTks, tkslist)

    def freq(self, tk):
        k = self.key_(tk)
        if k not in self.trie_:
            return 0
        return int(math.exp(self.trie_[k][0]) * self.DENOMINATOR + 0.5)

    def tag(self, tk):
        k = self.key_(tk)
        if k not in self.trie_:
            return ""
        return self.trie_[k][1]

    def score_(self, tfts):
        B = 30
        F, L, tks = 0, 0, []
        for tk, (freq, tag) in tfts:
            F += freq
            L += 0 if len(tk) < 2 else 1
            tks.append(tk)
        F /= len(tks)
        L /= len(tks)
        if self.DEBUG:
            print("[SC]", tks, len(tks), L, F, B / len(tks) + L + F)
        return tks, B / len(tks) + L + F

    def sortTks_(self, tkslist):
        res = []
        for tfts in tkslist:
            tks, s = self.score_(tfts)
            res.append((tks, s))
        return sorted(res, key=lambda x: x[1], reverse=True)

    def merge_(self, tks):
        patts = [
            (r"[ ]+", " "),
            (r"([0-9\+\.,%\*=-]) ([0-9\+\.,%\*=-])", r"\1\2"),
        ]
        # for p,s in patts: tks = re.sub(p, s, tks)

        # if split chars is part of token
        res = []
        tks = re.sub(r"[ ]+", " ", tks).split(" ")
        s = 0
        while True:
            if s >= len(tks):
                break
            E = s + 1
            for e in range(s + 2, min(len(tks) + 2, s + 6)):
                tk = "".join(tks[s:e])
                if re.search(self.SPLIT_CHAR, tk) and self.freq(tk):
                    E = e
            res.append("".join(tks[s:E]))
            s = E

        return " ".join(res)

    def maxForward_(self, line):
        res = []
        s = 0
        while s < len(line):
            e = s + 1
            t = line[s:e]
            while e < len(line) and self.trie_.has_keys_with_prefix(
                    self.key_(t)):
                e += 1
                t = line[s:e]

            while e - 1 > s and self.key_(t) not in self.trie_:
                e -= 1
                t = line[s:e]

            if self.key_(t) in self.trie_:
                res.append((t, self.trie_[self.key_(t)]))
            else:
                res.append((t, (0, '')))

            s = e

        return self.score_(res)

    def maxBackward_(self, line):
        res = []
        s = len(line) - 1
        while s >= 0:
            e = s + 1
            t = line[s:e]
            while s > 0 and self.trie_.has_keys_with_prefix(self.rkey_(t)):
                s -= 1
                t = line[s:e]

            while s + 1 < e and self.key_(t) not in self.trie_:
                s += 1
                t = line[s:e]

            if self.key_(t) in self.trie_:
                res.append((t, self.trie_[self.key_(t)]))
            else:
                res.append((t, (0, '')))

            s -= 1

        return self.score_(res[::-1])

    def english_normalize_(self, tks):
        return [self.stemmer.stem(self.lemmatizer.lemmatize(t)) if re.match(r"[a-zA-Z_-]+$", t) else t for t in tks]

    def tokenize(self, line):
        line = self._strQ2B(line).lower()
        line = self._tradi2simp(line)
        zh_num = len([1 for c in line if is_chinese(c)])
        if zh_num == 0:
            return " ".join([self.stemmer.stem(self.lemmatizer.lemmatize(t)) for t in word_tokenize(line)])

        arr = re.split(self.SPLIT_CHAR, line)
        res = []
        for L in arr:
            if len(L) < 2 or re.match(
                    r"[a-z\.-]+$", L) or re.match(r"[0-9\.-]+$", L):
                res.append(L)
                continue
            # print(L)

            # use maxforward for the first time
            tks, s = self.maxForward_(L)
            tks1, s1 = self.maxBackward_(L)
            if self.DEBUG:
                print("[FW]", tks, s)
                print("[BW]", tks1, s1)

            diff = [0 for _ in range(max(len(tks1), len(tks)))]
            for i in range(min(len(tks1), len(tks))):
                if tks[i] != tks1[i]:
                    diff[i] = 1

            if s1 > s:
                tks = tks1

            i = 0
            while i < len(tks):
                s = i
                while s < len(tks) and diff[s] == 0:
                    s += 1
                if s == len(tks):
                    res.append(" ".join(tks[i:]))
                    break
                if s > i:
                    res.append(" ".join(tks[i:s]))

                e = s
                while e < len(tks) and e - s < 5 and diff[e] == 1:
                    e += 1

                tkslist = []
                self.dfs_("".join(tks[s:e + 1]), 0, [], tkslist)
                res.append(" ".join(self.sortTks_(tkslist)[0][0]))

                i = e + 1

        res = " ".join(self.english_normalize_(res))
        if self.DEBUG:
            print("[TKS]", self.merge_(res))
        return self.merge_(res)

    def fine_grained_tokenize(self, tks):
        tks = tks.split(" ")
        zh_num = len([1 for c in tks if c and is_chinese(c[0])])
        if zh_num < len(tks) * 0.2:
            res = []
            for tk in tks:
                res.extend(tk.split("/"))
            return " ".join(res)

        res = []
        for tk in tks:
            if len(tk) < 3 or re.match(r"[0-9,\.-]+$", tk):
                res.append(tk)
                continue
            tkslist = []
            if len(tk) > 10:
                tkslist.append(tk)
            else:
                self.dfs_(tk, 0, [], tkslist)
            if len(tkslist) < 2:
                res.append(tk)
                continue
            stk = self.sortTks_(tkslist)[1][0]
            if len(stk) == len(tk):
                stk = tk
            else:
                if re.match(r"[a-z\.-]+$", tk):
                    for t in stk:
                        if len(t) < 3:
                            stk = tk
                            break
                    else:
                        stk = " ".join(stk)
                else:
                    stk = " ".join(stk)

            res.append(stk)

        return " ".join(self.english_normalize_(res))


class RAGFlowDocxParser:

    def __extract_table_content(self, tb):
        df = []
        for row in tb.rows:
            df.append([c.text for c in row.cells])
        return self.__compose_table_content(pd.DataFrame(df))

    def __compose_table_content(self, df):

        def blockType(b):
            patt = [
                ("^(20|19)[0-9]{2}[年/-][0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^(20|19)[0-9]{2}年$", "Dt"),
                (r"^(20|19)[0-9]{2}[年/-][0-9]{1,2}月*$", "Dt"),
                ("^[0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^第*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}年*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}[ABCDE]$", "DT"),
                ("^[0-9.,+%/ -]+$", "Nu"),
                (r"^[0-9A-Z/\._~-]+$", "Ca"),
                (r"^[A-Z]*[a-z' -]+$", "En"),
                (r"^[0-9.,+-]+[0-9A-Za-z/$￥%<>（）()' -]+$", "NE"),
                (r"^.{1}$", "Sg")
            ]
            for p, n in patt:
                if re.search(p, b):
                    return n
            tks = [t for t in rag_tokenizer.tokenize(b).split(" ") if len(t) > 1]   # 对句子进行切分，和jieba类似。
            if len(tks) > 3:
                if len(tks) < 12:
                    return "Tx"
                else:
                    return "Lx"

            if len(tks) == 1 and rag_tokenizer.tag(tks[0]) == "nr":
                return "Nr"

            return "Ot"

        if len(df) < 2:
            return []
        max_type = Counter([blockType(str(df.iloc[i, j])) for i in range(
            1, len(df)) for j in range(len(df.iloc[i, :]))])
        max_type = max(max_type.items(), key=lambda x: x[1])[0]

        colnm = len(df.iloc[0, :])
        hdrows = [0]  # header is not nessesarily appear in the first line
        if max_type == "Nu":
            for r in range(1, len(df)):
                tys = Counter([blockType(str(df.iloc[r, j]))
                              for j in range(len(df.iloc[r, :]))])
                tys = max(tys.items(), key=lambda x: x[1])[0]
                if tys != max_type:
                    hdrows.append(r)

        lines = []
        for i in range(1, len(df)):
            if i in hdrows:
                continue
            hr = [r - i for r in hdrows]
            hr = [r for r in hr if r < 0]
            t = len(hr) - 1
            while t > 0:
                if hr[t] - hr[t - 1] > 1:
                    hr = hr[t:]
                    break
                t -= 1
            headers = []
            for j in range(len(df.iloc[i, :])):
                t = []
                for h in hr:
                    x = str(df.iloc[i + h, j]).strip()
                    if x in t:
                        continue
                    t.append(x)
                t = ",".join(t)
                if t:
                    t += ": "
                headers.append(t)
            cells = []
            for j in range(len(df.iloc[i, :])):
                if not str(df.iloc[i, j]):
                    continue
                cells.append(headers[j] + str(df.iloc[i, j]))
            lines.append(";".join(cells))

        if colnm > 3:
            return lines
        return ["\n".join(lines)]

    def __call__(self, fnm, from_page=0, to_page=100000):
        self.doc = Document(fnm) if isinstance(
            fnm, str) else Document(BytesIO(fnm))
        pn = 0  # parsed page
        secs = []   # parsed contents
        for p in self.doc.paragraphs:
            if pn > to_page:
                break

            runs_within_single_paragraph = []   # save runs within the range of pages
            for run in p.runs:
                if pn > to_page:
                    break
                if from_page <= pn < to_page and p.text.strip():
                    runs_within_single_paragraph.append(run.text)   # append run.text first

                # wrap page break checker into a static method
                if RAGFlowDocxParser.has_page_break(run._element.xml):
                    pn += 1

            secs.append(("".join(runs_within_single_paragraph), p.style.name))  # then concat run.text as part of the paragraph

        tbls = [self.__extract_table_content(tb) for tb in self.doc.tables]
        return secs, tbls

class Docx(DocxParser):
    def __init__(self):
        pass

    def __clean(self, line):
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    def old_call(self, filename, binary=None, from_page=0, to_page=100000):
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            if from_page <= pn < to_page and p.text.strip():
                lines.append(self.__clean(p.text))
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        return [l for l in lines if l]

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        bull = bullets_category([p.text for p in self.doc.paragraphs])
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            question_level, p_text = docx_question_level(p, bull)
            if not p_text.strip("\n"):continue
            lines.append((question_level, p_text))

            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1

        visit = [False for _ in range(len(lines))]
        sections = []
        for s in range(len(lines)):
            e = s + 1
            while e < len(lines):
                if lines[e][0] <= lines[s][0]:
                    break
                e += 1
            if e - s == 1 and visit[s]: continue
            sec = []
            next_level = lines[s][0] + 1
            while not sec and next_level < 22:
                for i in range(s+1, e):
                    if lines[i][0] != next_level: continue
                    sec.append(lines[i][1])
                    visit[i] = True
                next_level += 1
            sec.insert(0, lines[s][1])

            sections.append("\n".join(sec))
        return [l for l in sections if l]

    def __str__(self) -> str:
        return f'''
            question:{self.question},
            answer:{self.answer},
            level:{self.level},
            childs:{self.childs}
        '''

def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Chinese", callback=None, **kwargs):
    """
        Supported file formats are docx, pdf, txt.
    """
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    pdf_parser = None
    sections = []
    # is it English
    eng = lang.lower() == "english"  # is_english(sections)

    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        for txt in Docx()(filename, binary):
            sections.append(txt)
        callback(0.8, "Finish parsing.")
        chunks = sections
        return tokenize_chunks(chunks, doc, eng, pdf_parser)
